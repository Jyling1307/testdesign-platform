"""Parse .docx files using python-docx."""


def parse_docx(file_path):
    """Extract text and structure from a .docx file.

    Returns:
        tuple: (raw_text, parsed_structure)
    """
    from docx import Document

    doc = Document(file_path)
    sections = []
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else 'Normal'
        level = 0
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip() or '0')
            except ValueError:
                level = 0

        section = {
            'text': text,
            'style': style_name,
            'level': level,
        }
        sections.append(section)
        full_text_parts.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                sections.append({
                    'text': row_text,
                    'style': 'Table',
                    'level': 0,
                })
                full_text_parts.append(row_text)

    raw_text = '\n'.join(full_text_parts)
    return raw_text, {'sections': sections}
